import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_project_docx():
    doc = Document()
    
    # Define standard styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- Title Page ---
    for _ in range(5): doc.add_paragraph()
    
    title = doc.add_heading('PROJECT REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Network Scanner Tool', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run('\nReal-time Network Monitoring and Security Auditing System\n')
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(3): doc.add_paragraph()
    
    p = doc.add_paragraph('Submitted as a Final Year Engineering Project')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # --- Abstract ---
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        "In the modern digital era, network security remains a paramount concern. This project presents a "
        "Real-time Network Monitoring and Security Auditing System, a comprehensive tool designed to provide "
        "deep visibility into network activities. By utilizing low-level packet manipulation via Scapy and "
        "a high-level management interface built on Django, the system discovers active hosts, identifies "
        "open ports, and detects device manufacturers. The result is a robust, scalable, and user-friendly "
        "platform that simplifies network auditing and strengthens security posture."
    )
    
    # --- 1. Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "Network scanning is the process of identifying active hosts on a network for inventory and security "
        "monitoring. This project aims to bridge the gap between complex command-line tools and modern visual dashboards."
    )
    
    doc.add_heading('1.1 Problem Statement', level=2)
    doc.add_paragraph(
        "Existing open-source tools often require advanced technical knowledge to operate and lack integrated "
        "reporting systems. There is a need for a unified platform that offers automatic discovery, visual health "
        "representation, and cross-platform accessibility via mobile synchronization."
    )
    
    doc.add_heading('1.2 Objectives', level=2)
    objectives = [
        'Automate identification of active hosts using ARP requests.',
        'Perform fast, multi-threaded TCP port scanning.',
        'Implement a live-updating dashboard with high-end aesthetics.',
        'Enable instant mobile monitoring via QR code synchronization.',
        'Facilitate professional documentation and reporting.'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
        
    # --- 2. Methodology & Algorithm ---
    doc.add_heading('2. Methodology & Algorithm', level=1)
    doc.add_heading('2.1 ARP Discovery', level=2)
    doc.add_paragraph(
        "The system utilizes the Address Resolution Protocol (ARP) for local discovery. The algorithm constructs "
        "an ARP request, broadcasts it to the subnet, and parses response packets to uniquely identify devices "
        "by their MAC and IP addresses."
    )
    
    doc.add_heading('2.2 Parallel Port Scanning', level=2)
    doc.add_paragraph(
        "To optimize performance, the system implements a multi-threaded approach using Python’s ThreadPoolExecutor. "
        "This allows the scanner to check thousands of ports per minute by running multiple connection attempts "
        "in parallel."
    )
    
    # --- 3. Technology Stack ---
    doc.add_heading('3. Technology Stack', level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology'
    hdr_cells[2].text = 'Rationale'
    
    tech_data = [
        ('Backend', 'Python', 'Versatile with strong networking libraries.'),
        ('Networking', 'Scapy, Socket', 'Raw packet control and TCP communication.'),
        ('Web UI', 'Django / JS', 'Secure framework and dynamic user interface.'),
        ('Database', 'SQLite', 'Lightweight and portable data storage.'),
        ('Reporting', 'ReportLab', 'Pixel-perfect PDF report generation.')
    ]
    for comp, tech, rat in tech_data:
        row_cells = table.add_row().cells
        row_cells[0].text = comp
        row_cells[1].text = tech
        row_cells[2].text = rat
    table.style = 'Table Grid'
    
    # --- 4. System Architecture ---
    doc.add_page_break()
    doc.add_heading('4. System Architecture', level=1)
    doc.add_paragraph(
        "The system follows a modular architecture where the scanning engine is decoupled from the user interface "
        "to ensure high performance and maintainability."
    )
    
    img_path = 'documentation/flowchart_viva.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        doc.add_paragraph('Figure 1: High-level System Architecture and Data Flow', style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 5. Implementation Samples ---
    doc.add_heading('5. Core Implementation', level=1)
    doc.add_heading('ARP Scanner Logic', level=2)
    code = (
        "def scan_network(ip_range):\n"
        "    arp_request = scapy.ARP(pdst=ip_range)\n"
        "    broadcast = scapy.Ether(dst='ff:ff:ff:ff:ff:ff')\n"
        "    answered_list = scapy.srp(broadcast / arp_request, timeout=1, verbose=False)[0]\n"
        "    return [{'ip': e[1].psrc, 'mac': e[1].hwsrc} for e in answered_list]"
    )
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    # --- 6. Future Scope ---
    doc.add_heading('6. Future Scope', level=1)
    doc.add_paragraph(
        "Future enhancements include AI-driven intrusion detection, integration with CVE vulnerability "
        "databases for live security alerting, and cloud-based multi-network synchronization."
    )
    
    # --- Conclusion ---
    doc.add_heading('7. Conclusion', level=1)
    doc.add_paragraph(
        "This project successfully bridges the gap between complex networking tools and professional user "
        "experiences. It provides a robust platform for real-time security auditing, fulfilling all academic "
        "and technical requirements for an engineering project."
    )
    
    # --- Save ---
    output_path = 'documentation/Project_Report_Network_Scanner.docx'
    doc.save(output_path)
    print(f"Professional Word document updated at: {output_path}")

if __name__ == "__main__":
    generate_project_docx()
